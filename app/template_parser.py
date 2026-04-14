import os
import shutil
from pathlib import Path
from docx import Document as DocxDocument
from docx.oxml.ns import qn

DATA_DIR = os.environ.get("DATA_DIR", "../data")
TEMPLATES_DIR = os.path.join(DATA_DIR, "templates")


def save_template_file(filename: str, content: bytes) -> str:
    """Guarda el fichero de plantilla en disco y devuelve la ruta."""
    os.makedirs(TEMPLATES_DIR, exist_ok=True)
    dest = os.path.join(TEMPLATES_DIR, filename)
    with open(dest, "wb") as f:
        f.write(content)
    return dest


def delete_template_file(filename: str):
    path = os.path.join(TEMPLATES_DIR, filename)
    if os.path.exists(path):
        os.remove(path)


def extract_sections_docx(filename: str) -> list[str]:
    """
    Extrae los títulos de sección (Heading 1) de un .docx.
    Devuelve lista de strings con los títulos encontrados.
    """
    path = os.path.join(TEMPLATES_DIR, filename)
    doc = DocxDocument(path)
    sections = []
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading 1") or para.style.name.startswith("Título 1"):
            text = para.text.strip()
            if text:
                sections.append(text)
    # Si no hay Heading 1, intentar con Heading 2
    if not sections:
        for para in doc.paragraphs:
            if para.style.name.startswith("Heading") or para.style.name.startswith("Título"):
                text = para.text.strip()
                if text:
                    sections.append(text)
    return sections


def extract_sections_odt(filename: str) -> list[str]:
    """
    Extrae los títulos de sección de un .odt.
    """
    try:
        from odf.opendocument import load as odf_load
        from odf.text import H
        from odf.element import Element

        path = os.path.join(TEMPLATES_DIR, filename)
        doc = odf_load(path)
        sections = []

        def get_text(node) -> str:
            result = ""
            if hasattr(node, "childNodes"):
                for child in node.childNodes:
                    if child.nodeType == child.TEXT_NODE:
                        result += child.data
                    else:
                        result += get_text(child)
            return result

        for elem in doc.text.childNodes:
            if elem.__class__.__name__ == "H":
                outline = elem.getAttribute("outlinelevel")
                if outline in ("1", None):
                    text = get_text(elem).strip()
                    if text:
                        sections.append(text)
        return sections
    except Exception as e:
        return []


def extract_sections(filename: str, file_type: str) -> list[str]:
    if file_type == "docx":
        return extract_sections_docx(filename)
    elif file_type == "odt":
        return extract_sections_odt(filename)
    return []


def get_template_path(filename: str) -> str:
    return os.path.join(TEMPLATES_DIR, filename)
