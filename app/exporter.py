import os
import shutil
import subprocess
from pathlib import Path
from docx import Document as DocxDocument
from models import Document

DATA_DIR = os.environ.get("DATA_DIR", "../data")
OUTPUTS_DIR = os.path.join(DATA_DIR, "outputs")
TEMPLATES_DIR = os.path.join(DATA_DIR, "templates")


def _get_output_path(project_id: str, ext: str) -> str:
    os.makedirs(OUTPUTS_DIR, exist_ok=True)
    return os.path.join(OUTPUTS_DIR, f"{project_id}.{ext}")


def _detect_styles(doc: DocxDocument) -> tuple[str, str]:
    """
    Detecta el nombre del estilo de título y cuerpo disponibles en la plantilla.
    Devuelve (heading_style, body_style).
    """
    style_names = {s.name for s in doc.styles}

    # Buscar estilo de título (por prioridad)
    heading_candidates = [
        "Heading 1", "Título 1", "heading 1", "titulo 1",
        "Heading1", "Titulo1", "Header 1",
    ]
    heading_style = next((s for s in heading_candidates if s in style_names), "Heading 1")

    # Buscar estilo de cuerpo
    body_candidates = [
        "Normal", "Body Text", "Cuerpo de texto", "Default Paragraph Style",
        "Text Body", "Texto independiente",
    ]
    body_style = next((s for s in body_candidates if s in style_names), "Normal")

    return heading_style, body_style



def _add_content_paragraph(doc: DocxDocument, text: str, style: str):
    """Añade un párrafo de contenido respetando listas con guión."""
    if text.startswith("- ") or text.startswith("• "):
        # Intentar usar estilo de lista si existe
        list_styles = {"List Bullet", "Lista con viñetas", "List Paragraph"}
        available = {s.name for s in doc.styles}
        list_style = next((s for s in list_styles if s in available), style)
        p = doc.add_paragraph(text.lstrip("- •").strip(), style=list_style)
    else:
        doc.add_paragraph(text, style=style)


def export_docx(document: Document, template_filename: str, project_name: str) -> str:
    """
    Genera un .docx usando la plantilla como base.
    Preserva portada, índice, encabezado y pie de página intactos
    y añade el contenido generado a continuación.
    """
    template_path = os.path.join(TEMPLATES_DIR, template_filename)
    output_path = _get_output_path(document.project_id, "docx")

    shutil.copy2(template_path, output_path)
    doc = DocxDocument(output_path)

    heading_style, body_style = _detect_styles(doc)

    sections = [
        s for s in sorted(document.sections, key=lambda s: s.order)
        if s.status.value in ("generated", "approved")
    ]

    if not sections:
        doc.save(output_path)
        return output_path

    # Salto de página para separar la plantilla del contenido generado
    doc.add_page_break()

    for section in sections:
        doc.add_paragraph(section.title, style=heading_style)

        if section.content.strip():
            for line in section.content.split("\n"):
                line = line.strip()
                if line:
                    _add_content_paragraph(doc, line, body_style)

    doc.save(output_path)
    return output_path


def export_odt(document: Document, template_filename: str, project_name: str) -> str:
    docx_path = export_docx(document, template_filename, project_name)
    output_path = _get_output_path(document.project_id, "odt")
    _convert_with_libreoffice(docx_path, output_path, "odt")
    return output_path


def export_pdf(document: Document, template_filename: str, project_name: str) -> str:
    docx_path = export_docx(document, template_filename, project_name)
    output_path = _get_output_path(document.project_id, "pdf")
    _convert_with_libreoffice(docx_path, output_path, "pdf")
    return output_path


def _convert_with_libreoffice(input_path: str, output_path: str, fmt: str):
    out_dir = os.path.dirname(output_path)
    result = subprocess.run(
        ["libreoffice", "--headless", "--convert-to", fmt, "--outdir", out_dir, input_path],
        capture_output=True,
        text=True,
        timeout=60,
    )
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice error: {result.stderr}")

    generated = os.path.join(out_dir, Path(input_path).stem + f".{fmt}")
    if generated != output_path and os.path.exists(generated):
        shutil.move(generated, output_path)
